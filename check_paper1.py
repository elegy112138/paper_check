from config import get_collection
from annotate import add_comments
from llm import *
from tqdm import tqdm
from concurrent.futures import ThreadPoolExecutor, as_completed
import block
from docx import Document
import process
from bson.objectid import ObjectId

# 假设有一个全局字典来跟踪任务状态


url_base_list = [
    "https://api.rcouyi.com/v1/",
    "https://us.rcouyi.com/v1/",
    "https://us-1.rcouyi.com/v1/",
    "https://jp-5.rcouyi.com/v1/"
]

def process_file(file_data, task_id):
    collection = get_collection("paper")
    process.tasks[task_id] = {'total': 0, 'current': 0, 'status': '未开始',"id":''}
    for file in file_data:
        process.tasks[task_id]['current'] =0
        collection.update_one({"_id": ObjectId(file.get('id'))}, {"$set": {"status": 0}})
        process.tasks[task_id]['id']=file.get('id')
        try:
            # 假设这是调用check_paper函数的代码
            if process.cancel:
                pass
            else:
                check_paper(file.get('file_path'), file.get('unique_filename'), task_id)
        except Exception as e:
            # 当check_paper函数抛出任何异常时，这里会捕获到
            print("捕获到异常：", str(e))
            collection.update_one({"_id":ObjectId(file.get('id'))},{"$set":{"status":3}})
            continue
        # check_paper(file.get('file_path'), file.get('unique_filename'), task_id)
        if process.cancel:
            print("执行取消操作")
            collection.update_one({"_id": ObjectId(file.get('id'))}, {"$set": {"status": 4}})
            continue
        collection.update_one({"_id":ObjectId(file.get('id'))},{"$set":{"status":2}})
    process.tasks.pop(task_id)




def check_paper(file_path, unique_filename, task_id):
    if process.cancel:
        print(1111)
        return
    doc = Document(file_path)
    key_sections = block.find_key_sections(doc)
    process.tasks[task_id]['total'] = key_sections.get('Acknowledgements')
    process.tasks[task_id]['status']='进行中'
    processed_filename = f"processed_{unique_filename}"
    paper_comments = []
    paper_comments = check_front(doc, key_sections['Abstract'], paper_comments, task_id)
    paper_comments, doc = check_abstract(doc, key_sections['Abstract'], key_sections['TextStart'], paper_comments, task_id)
    paper_comments, doc = check_text(doc, key_sections['TextStart'], key_sections['References'], paper_comments, task_id)
    paper_comments = check_references(doc, key_sections['References'],
                                      key_sections['Appendix'] if key_sections['Appendix'] is not None else
                                      key_sections['Acknowledgements'], paper_comments, task_id)
    processed_filepath = os.path.join('processed/', processed_filename)

    print(paper_comments)
    doc.save(processed_filepath)
    add_comments(paper_comments, processed_filepath)
    return processed_filename


def check_front(doc, pos, paper_comments, task_id):
    keywords = ['本科毕业论文（设计）']  # 包含所有要检查的关键词
    for i in tqdm(range(pos), desc="封面"):
        if process.cancel:
            return []
        process.tasks[task_id]['current'] +=1
        para = doc.paragraphs[i]
        text = para.text.strip()
        # 检查段落文本是否以任一关键词开头
        if any(text.startswith(kw) for kw in keywords):
            paper_comments.append({'text': text,
                                   'comment1': '这一段的下面的内容应该使用华文仿宋，并且字号为三号，下面应有下划线\n(若已经正确则忽略这条批注)'})
    return paper_comments


def check_abstract(doc, start_pos, end_pos, paper_comments, task_id):
    flag = None
    for i in tqdm(range(start_pos, end_pos), desc="摘要"):
        if process.cancel:
            return [],doc
        process.tasks[task_id]['current'] +=1
        para = doc.paragraphs[i]
        text = para.text.strip()
        if len(text) > 30:
            advice1 = get_correction_suggestion(text, True)
            paper_comments.append({'text': text,
                                   'comment1': advice1})
        if text == '摘要':
            flag = True
            continue
        elif text == 'Abstract':
            flag = False
            continue
        if flag is not None:
            # 假设小四号字体大小约为12磅（Pt）
            if flag:  # 如果是中文摘要
                # advice2 = check_abstract_format(para, '宋体', 12)
                doc.paragraphs[i] = check_abstract_format(para, '宋体', 12)
            else:  # 如果是英文摘要
                doc.paragraphs[i] = check_abstract_format(para, 'Time New Roman', 12)

    return paper_comments, doc


def process_paragraph(para, index):
    if process.cancel:
        return None, index, None  # 没有建议时返回None
    api_url = url_base_list[index % 4]
    text = para.text.strip()
    level = 0
    style_name = para.style.name

    # 检查是否为标题，并获取级别
    if style_name == 'Heading 1':
        level = 1
    elif style_name == 'Heading 2':
        level = 2
    elif style_name == 'Heading 3':
        level = 3

    # 对标题或较长的文本段落进行处理
    if level != 0 or len(text) > 30:
        # 假设对段落进行的处理会直接修改para对象
        # 对于标题，调整其格式
        if level != 0:
            para = check_title_format(para, level)
            return para, index, None
        # 对于较长的段落，获取修改建议
        else:
            para = check_text_format(para)
            advice = get_correction_suggestion(text, True, api_url)
            return para, index, advice  # 假设这里只返回建议，而段落已被直接修改
    return None, index, None  # 没有建议时返回None


def check_text(doc, start_pos, end_pos, paper_comments, task_id):
    # 原始范围内的段落总数
    total_paragraphs = end_pos - start_pos

    # 使用列表推导式过滤并计算非空的段落
    paragraphs_to_process = [
        {'para': doc.paragraphs[i], 'index': i}
        for i in range(start_pos, end_pos)
        if doc.paragraphs[i].text.strip()  # 仅包含文本非空的段落
    ]

    # 计算被去除（即过滤掉的）段落的总数
    removed_paragraphs_count = total_paragraphs - len(paragraphs_to_process)
    process.tasks[task_id]['current'] += removed_paragraphs_count
    with ThreadPoolExecutor(max_workers=4) as executor:
        # 提交处理任务，同时传递段落和它的原始索引
        futures = [executor.submit(process_paragraph, item['para'], item['index']) for item in paragraphs_to_process]

        # 等待任务完成并收集结果
        for future in tqdm(as_completed(futures), total=len(paragraphs_to_process), desc="正文处理中"):
            process.tasks[task_id]['current'] += 1
            para, index, advice = future.result()
            if para:
                doc.paragraphs[index] = para
            # 使用索引和建议更新paper_comments
            if advice:
                paper_comments.append({'text': doc.paragraphs[index].text, 'comment1': advice})

    # 注意：这里不需要对doc.paragraphs进行重新排序或修改，因为每个段落已经在原地被修改
    return paper_comments, doc


def check_references(doc, start_pos, end_pos, paper_comments,task_id):
    # 原始范围内的段落总数
    total_paragraphs = end_pos - start_pos

    # 使用列表推导式过滤并计算非空的段落
    paragraphs_to_process = [
        {'para': doc.paragraphs[i], 'index': i}
        for i in range(start_pos, end_pos)
        if doc.paragraphs[i].text.strip()  # 仅包含文本非空的段落
    ]

    # 计算被去除（即过滤掉的）段落的总数
    removed_paragraphs_count = total_paragraphs - len(paragraphs_to_process)
    process.tasks[task_id]['current'] += removed_paragraphs_count

    with ThreadPoolExecutor(max_workers=4) as executor:
        # 提交处理任务，同时传递段落和它的原始索引
        futures = [executor.submit(process_references, item['para'], item['index']) for item in paragraphs_to_process]

        # 等待任务完成并收集结果
        for future in tqdm(as_completed(futures), total=len(paragraphs_to_process), desc="参考文献处理中"):
            process.tasks[task_id]['current'] += 1
            index, advice = future.result()
            if advice:
                paper_comments.append({'text': doc.paragraphs[index].text, 'comment1': advice})

    return paper_comments


def process_references(para, index):
    if process.cancel:
        return index, None
    api_url = url_base_list[index % 4]
    text = para.text
    if len(text) > 20:
        advice = get_correction_suggestion(text, False, api_url)
        return index, advice
    else:
        return index, None
